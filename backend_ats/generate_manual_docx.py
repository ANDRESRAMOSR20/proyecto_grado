from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path


def build_document() -> Document:
    doc = Document()

    # Title
    title = doc.add_paragraph()
    run = title.add_run("Manual de usuario – Plataforma de gestión de vacantes y preselección")
    run.bold = True
    run.font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    sections = [
        ("Propósito", (
            "Guiar a cualquier usuario (sin experiencia técnica) en el uso de la plataforma para "
            "buscar vacantes, postularse y seguir su proceso de selección. Incluye además el uso "
            "administrativo para publicar vacantes, gestionar etapas y revisar métricas."
        )),
        ("Alcance", (
            "Acceso al sistema, registro e inicio de sesión, actualización de perfil (metadatos y CV), "
            "consulta de ofertas, postulación, seguimiento por etapas; funciones de administrador (publicación, "
            "preselección y métricas)."
        )),
        ("1. Requisitos y acceso", (
            "Requisitos: navegador actualizado (Chrome, Edge o Firefox), conexión a internet estable y cuenta de usuario. "
            "Acceso: ingrese a la URL oficial de la plataforma. Si no tiene cuenta, seleccione 'Registrarse'."
        )),
        ("2. Registro de cuenta", (
            "Abra 'Registrarse', complete nombre, correo, documento de identidad y contraseña. Envíe el formulario; "
            "si no existen duplicados, su cuenta quedará creada. Luego, inicie sesión."
        )),
        ("3. Inicio de sesión", (
            "Ingrese correo y contraseña en 'Iniciar sesión'. Opción 'Recordarme': activada mantiene la sesión; desactivada, "
            "se cierra al salir del navegador. Si hay error, verifique datos o consulte soporte."
        )),
        ("4. Perfil del usuario (metadatos y CV)", (
            "En 'Perfil', complete/actualice: nombre completo, número de celular y cargue su CV en PDF. Guarde los cambios. "
            "El CV vinculado mejora la preselección comparando su contenido con el 'perfil ideal' de cada vacante."
        )),
        ("5. Ver ofertas", (
            "En 'Ofertas', busque por palabras clave (título o descripción), abra una vacante para ver detalles."
        )),
        ("6. Postularse a una vacante", (
            "En la vacante abierta, seleccione 'Postular ahora'. La plataforma registra su postulación y aparece en 'Selección'. "
            "Mantenga el perfil completo (incluido CV) para mejores coincidencias."
        )),
        ("7. Ver selección (seguimiento de postulaciones)", (
            "En 'Selección' verá: estado general (En proceso, Finalizado, Descartado) y línea de tiempo (postulación, entrevista, "
            "prueba, resultado) con estado, fecha y feedback (enlaces clicables si aplica). Puede eliminar una postulación."
        )),
        ("8. Cierre de sesión", (
            "Use 'Cerrar sesión' para finalizar de forma segura, especialmente en equipos compartidos."
        )),
        ("9. Funciones para administradores", (
            "Acceda al panel solo con permisos de administrador."
        )),
        ("9.1 Publicar vacantes", (
            "Cree vacantes con título (obligatorio), descripción (obligatoria), 'perfil ideal' (opcional) y fecha de publicación "
            "(opcional). Edite/Elimine desde la lista y guarde cambios."
        )),
        ("9.2 Sistema de preselección", (
            "Vea todas las postulaciones con datos del candidato, vacante, estado general, progreso por etapas y porcentaje de similitud "
            "(si hay 'perfil ideal'). Busque, filtre y ordene. Realice cambios masivos y registre feedback/enlaces desde el modal. "
            "Actualizar 'resultado' a aceptado/rechazado ajusta el estado general."
        )),
        ("9.3 Métricas", (
            "Consulte totales (usuarios, vacantes, postulaciones) y gráficas por estado, por vacante y series temporales."
        )),
        ("10. Buenas prácticas y privacidad", (
            "Mantenga sus credenciales seguras, cargue CVs en PDF legibles, mantenga su perfil al día. La plataforma aplica "
            "controles de acceso por rol y tratamiento responsable de datos."
        )),
        ("11. Resolución de problemas", (
            "Inicio de sesión: verifique credenciales y conexión; reintente o contacte soporte. Carga de CV: confirme PDF y reintente. "
            "No aparecen ofertas/postulaciones: recargue o cierre/abra sesión. Cambios de etapas (admin): actualice, verifique permisos."
        )),
        ("12. Soporte", (
            "Contacto institucional y horario según implementación. Para soporte, envíe descripción del problema, capturas y fecha/hora."
        )),
        ("13. Entregables del proyecto (académico)", (
            "Repositorio privado (GitHub) con acceso a evaluadores y equipo; incluye instrucciones de instalación/ejecución y despliegue. "
            "Video demostrativo (5 minutos): objetivos; registro/login; perfil y CV; ofertas y postulación; 'Selección' con timeline; "
            "administración (publicar y preselección); métricas; cierre con beneficios y próximos pasos."
        )),
        ("14. Glosario", (
            "Perfil ideal: texto con competencias esperadas para comparar con el CV. Timeline: etapas con estado, fecha y feedback. "
            "Preselección: priorización de candidatos apoyada en similitud semántica. Similitud: % de ajuste entre perfil ideal y CV."
        )),
    ]

    for heading, text in sections:
        p = doc.add_paragraph()
        r = p.add_run(heading)
        r.bold = True
        r.font.size = Pt(12)
        body = doc.add_paragraph(text)
        body.paragraph_format.space_after = Pt(8)

    return doc


def main():
    out_path = Path(r"d:\webstorm\docs\Manual_de_Usuario.docx")
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(out_path)
    print(f"Documento generado: {out_path}")


if __name__ == "__main__":
    main()



